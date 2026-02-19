"""
AI Prompt Generator for AI Exercises
Generates Daily Prompts and Friday Fun Prompts
Supports Claude AI, GitHub Copilot, ChatGPT, Gemini
Restricted to @ssaandco.com and authorized domain email addresses
"""

import streamlit as st
import json
import zipfile
import io
import re
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load .env file if present (SSA default keys live there, never in source code)
try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).parent / ".env", override=False)
except ImportError:
    pass  # python-dotenv not installed — keys must be entered manually

# ─── Page Config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="AI Prompt Generator",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

SSA_ADMIN_DOMAIN = "@ssaandco.com"
DOMAINS_FILE     = Path(__file__).parent / "data_storage" / "authorized_domains.json"

def _load_domains() -> list:
    try:
        if DOMAINS_FILE.exists():
            data = json.loads(DOMAINS_FILE.read_text(encoding="utf-8"))
            domains = data.get("domains", [])
            if isinstance(domains, list) and domains:
                if SSA_ADMIN_DOMAIN not in domains:
                    domains.insert(0, SSA_ADMIN_DOMAIN)
                return domains
    except Exception:
        pass
    return [SSA_ADMIN_DOMAIN]

def _save_domains(domains: list) -> None:
    if SSA_ADMIN_DOMAIN not in domains:
        domains.insert(0, SSA_ADMIN_DOMAIN)
    DOMAINS_FILE.parent.mkdir(parents=True, exist_ok=True)
    DOMAINS_FILE.write_text(
        json.dumps({"domains": domains}, indent=2), encoding="utf-8"
    )

if "allowed_domains" not in st.session_state:
    st.session_state["allowed_domains"] = _load_domains()

ALLOWED_DOMAINS = st.session_state["allowed_domains"]

INDUSTRIES = [
    "",
    "Industry-Agnostic",
    "Healthcare",
    "Finance & Banking",
    "Legal",
    "Technology",
    "Manufacturing",
    "Retail & E-Commerce",
    "Education",
    "Government & Public Sector",
    "Real Estate",
    "Consulting & Professional Services",
    "Marketing & Advertising",
    "Logistics & Supply Chain",
    "Insurance",
    "Energy & Utilities",
    "Non-Profit",
]

AI_AGENTS = ["", "Claude AI", "GitHub Copilot", "ChatGPT", "Gemini"]

TEMPLATES_DIR = Path(__file__).parent / "templates"

# ─── Custom CSS ───────────────────────────────────────────────────────────────
st.markdown("""
<style>
    .stApp { background-color: #f0f4f8; }

    .main-header {
        background: linear-gradient(135deg, #002060 0%, #1a3a8f 100%);
        color: white; padding: 2rem 2.5rem; border-radius: 12px;
        margin-bottom: 1.5rem; box-shadow: 0 4px 20px rgba(0,32,96,0.3);
    }
    .main-header h1 { color: white; margin: 0; font-size: 2rem; }
    .main-header p  { color: rgba(255,255,255,0.85); margin: 0.5rem 0 0 0; font-size: 1.05rem; }

    .auth-card {
        background: white; border-radius: 14px; padding: 3rem 2.5rem;
        max-width: 480px; margin: 3rem auto;
        box-shadow: 0 8px 32px rgba(0,32,96,0.12); text-align: center;
    }
    .auth-card .lock-icon { font-size: 3.5rem; margin-bottom: 1rem; }
    .auth-card h2 { color: #002060; margin: 0 0 0.5rem 0; }
    .auth-card p  { color: #555; margin: 0 0 1.2rem 0; font-size: 0.95rem; }
    .auth-domains {
        background: #f0f4f8; border-radius: 8px; padding: 0.6rem 1rem;
        font-size: 0.85rem; color: #002060; font-weight: 600;
        margin-bottom: 1.5rem; display: inline-block;
    }
    .auth-error {
        background: #fff0f0; border: 1px solid #ffcccc; border-radius: 8px;
        padding: 0.7rem 1rem; color: #cc0000; font-size: 0.9rem; margin-top: 0.8rem;
    }

    .field-label {
        font-weight: 700; color: #002060; font-size: 0.8rem;
        text-transform: uppercase; letter-spacing: 0.05em; margin-bottom: 0.2rem;
    }
    .field-value {
        background: #f8f9fc; border-radius: 6px; padding: 0.6rem 0.8rem;
        font-size: 0.9rem; margin-bottom: 0.8rem; border: 1px solid #e8eef5; white-space: pre-wrap;
    }

    .badge-daily {
        background: #002060; color: white; padding: 0.2rem 0.7rem; border-radius: 20px;
        font-size: 0.75rem; font-weight: 600; display: inline-block; margin-bottom: 0.5rem;
    }
    .badge-friday {
        background: #7030a0; color: white; padding: 0.2rem 0.7rem; border-radius: 20px;
        font-size: 0.75rem; font-weight: 600; display: inline-block; margin-bottom: 0.5rem;
    }

    .status-bar {
        background: #002060; color: white; padding: 0.8rem 1.2rem;
        border-radius: 8px; margin-bottom: 1rem; font-size: 0.9rem;
    }

    .stButton > button {
        background: linear-gradient(135deg, #002060, #1a3a8f);
        color: white; border: none; border-radius: 8px;
        padding: 0.7rem 2rem; font-size: 1rem; font-weight: 600; width: 100%; transition: all 0.2s;
    }
    .stButton > button:hover {
        background: linear-gradient(135deg, #001540, #0d2870);
        transform: translateY(-1px); box-shadow: 0 4px 12px rgba(0,32,96,0.3);
    }

    .download-area {
        background: linear-gradient(135deg, #f0f7ff, #e8f0fe);
        border: 2px dashed #002060; border-radius: 10px;
        padding: 1.5rem; text-align: center; margin-top: 1rem;
    }

    .user-pill {
        background: #e8f0fe; border-radius: 20px; padding: 0.4rem 0.9rem;
        font-size: 0.8rem; color: #002060; font-weight: 600;
        margin-bottom: 1rem; display: inline-block; word-break: break-all;
    }

    .legend-bar {
        display: flex; align-items: center; gap: 1.5rem;
        background: white; border-radius: 8px; padding: 0.6rem 1.2rem;
        box-shadow: 0 1px 4px rgba(0,0,0,0.07); margin-bottom: 1rem;
        font-size: 0.85rem; color: #444; flex-wrap: wrap;
    }
    .legend-item { display: flex; align-items: center; gap: 6px; }
    .legend-dot {
        display: inline-block; width: 13px; height: 13px;
        border-radius: 50%; flex-shrink: 0;
    }

    .attachment-box {
        background: #fff8e1; border: 1px solid #ffe082; border-radius: 8px;
        padding: 0.75rem 0.9rem; margin-bottom: 0.8rem;
    }
    .att-title {
        font-weight: 700; color: #7a5c00; font-size: 0.78rem;
        text-transform: uppercase; letter-spacing: 0.05em; margin-bottom: 0.3rem;
    }
    .att-desc { font-size: 0.88rem; color: #5a4000; margin-bottom: 0.5rem; }
</style>
""", unsafe_allow_html=True)


# ─── Auth Helpers ─────────────────────────────────────────────────────────────

def is_valid_email(email: str) -> bool:
    return bool(re.match(r'^[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}$', email.strip()))

def is_allowed_domain(email: str) -> bool:
    domains = st.session_state.get("allowed_domains", ALLOWED_DOMAINS)
    return any(email.strip().lower().endswith(d.strip().lower()) for d in domains)

def is_ssa_admin(email: str) -> bool:
    return email.strip().lower().endswith(SSA_ADMIN_DOMAIN)

def render_auth_gate():
    st.markdown("""
    <div class="main-header">
        <h1>🤖 AI Prompt Generator</h1>
        <p>Generate Daily &amp; Friday Fun Prompts for AI Adoption Exercises</p>
    </div>
    <div class="auth-card">
        <div class="lock-icon">🔐</div>
        <h2>Restricted Access</h2>
        <p>This tool is exclusively available to <strong>SSA &amp; Company employees</strong>
        and <strong>authorized partner organizations</strong>.<br><br>
        Please sign in with your work email address to continue. If you believe you should
        have access and are unable to log in, please contact your SSA &amp; Company representative.</p>
    </div>
    """, unsafe_allow_html=True)

    _, col, _ = st.columns([1, 2, 1])
    with col:
        email_input = st.text_input(
            "Work email address",
            placeholder="name@yourcompany.com",
            key="auth_email_input"
        )
        if st.button("✅ Login", use_container_width=True):
            email = email_input.strip()
            if not email:
                st.markdown('<div class="auth-error">⚠️ Please enter your email address.</div>', unsafe_allow_html=True)
            elif not is_valid_email(email):
                st.markdown('<div class="auth-error">⚠️ That does not look like a valid email address.</div>', unsafe_allow_html=True)
            elif not is_allowed_domain(email):
                st.markdown(
                    '<div class="auth-error">🚫 Access is restricted to SSA &amp; Company employees and authorized partners. '
                    'If you require access, please contact your SSA &amp; Company representative.</div>',
                    unsafe_allow_html=True
                )
            else:
                st.session_state['authenticated'] = True
                st.session_state['user_email'] = email.lower()
                st.rerun()


# ─── Auth Gate ────────────────────────────────────────────────────────────────

if not st.session_state.get('authenticated'):
    render_auth_gate()
    st.stop()

user_email = st.session_state.get('user_email', '')

# ─── Authenticated Header ─────────────────────────────────────────────────────
st.markdown("""
<div class="main-header">
    <h1>🤖 AI Prompt Generator</h1>
    <p>Generate Daily &amp; Friday Fun Prompts for AI Adoption Exercises</p>
</div>
""", unsafe_allow_html=True)

# ─── Sidebar ──────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown(f'<div class="user-pill">👤 {user_email}</div>', unsafe_allow_html=True)
    if st.button("🔓 Sign Out", use_container_width=True):
        for k in ['authenticated', 'user_email', 'generated_prompts', 'is_friday_list']:
            st.session_state.pop(k, None)
        st.rerun()

    # ── Access Management (SSA admins only) ─────────────────────────────────
    if is_ssa_admin(user_email):
        with st.expander("🔑 Access Management", expanded=False):
            st.markdown(
                "<small style='color:#888;'>Manage which organizations can access this tool. "
                "Changes are <strong>saved permanently</strong> and survive app restarts.</small>",
                unsafe_allow_html=True
            )
            domains = st.session_state["allowed_domains"]

            # Show current list with remove buttons
            st.markdown("**Authorized domains:**")
            to_remove = None
            for d in domains:
                col_d, col_x = st.columns([5, 1])
                col_d.markdown(
                    f"<div style='background:#f0f4f8;border-radius:6px;padding:4px 10px;"
                    f"font-size:0.85rem;font-family:monospace;margin-bottom:4px;'>{d}</div>",
                    unsafe_allow_html=True
                )
                if col_x.button("✕", key=f"rm_{d}", help=f"Remove {d}",
                                 disabled=(d.strip().lower() == SSA_ADMIN_DOMAIN)):
                    to_remove = d
            if to_remove:
                updated = [d for d in st.session_state["allowed_domains"] if d != to_remove]
                st.session_state["allowed_domains"] = updated
                _save_domains(updated)
                st.rerun()

            st.markdown("**Add a new domain:**")
            new_dom_col, add_col = st.columns([3, 1])
            new_domain = new_dom_col.text_input(
                "Domain", placeholder="@partner.com",
                key="new_domain_input", label_visibility="collapsed"
            )
            if add_col.button("Add", key="add_domain_btn"):
                nd = new_domain.strip().lower()
                if nd and not nd.startswith("@"):
                    nd = "@" + nd
                if nd and nd not in st.session_state["allowed_domains"]:
                    if re.match(r'^@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}$', nd):
                        updated = st.session_state["allowed_domains"] + [nd]
                        st.session_state["allowed_domains"] = updated
                        _save_domains(updated)
                        st.success(f"✅ {nd} added and saved permanently.")
                        st.rerun()
                    else:
                        st.error("Invalid domain format.")
                elif nd in st.session_state["allowed_domains"]:
                    st.warning(f"{nd} is already authorized.")

            st.caption(
                f"🔒 SSA & Company domain ({SSA_ADMIN_DOMAIN}) is always authorized and cannot be removed."
            )

    st.markdown("---")
    st.markdown("## ⚙️ Configuration")

    # ── AI Agent ─────────────────────────────────────────────────────────────
    st.markdown("### 🤖 AI Agent")
    ai_agent = st.selectbox(
        "Select AI Agent",
        AI_AGENTS,
        index=0,
        format_func=lambda x: "— Select an AI Agent —" if x == "" else x,
        help="Choose which AI agent will generate and respond to the prompts."
    )

    # ── API Key ──────────────────────────────────────────────────────────────
    _key_hints = {
        "Claude AI":      ("sk-ant-...",               "Make sure to add a valid API key for this agent."),
        "ChatGPT":        ("sk-proj-...",               "Make sure to add a valid API key for this agent."),
        "Gemini":         ("AIza...",                   "Make sure to add a valid API key for this agent."),
        "GitHub Copilot": ("ghp_... or github_pat_...", "Make sure to add a valid API key for this agent."),
    }
    st.markdown("### 🔑 API Key")
    _hint = _key_hints.get(ai_agent, ("Your API key", ""))
    api_key = st.text_input(
        "Enter your API Key",
        type="password",
        placeholder=_hint[0],
        help=f"Required to call {ai_agent}. {_hint[1]}. Keys are never stored.",
        key="api_key_input"
    )
    if _hint[1]:
        st.caption(f"🔒 Never saved. 💡 {_hint[1]}")
    else:
        st.caption("🔒 Used only for this session and never saved.")

    st.markdown("---")

    # ── Industry ─────────────────────────────────────────────────────────────
    st.markdown("### 🏭 Industry")
    industry = st.selectbox(
        "Select Industry",
        INDUSTRIES,
        index=0,
        format_func=lambda x: "— Select an Industry —" if x == "" else x,
        help="Prompts will be tailored to this industry's workflows and terminology."
    )

    st.markdown("---")

    # ── Prompt Settings ───────────────────────────────────────────────────────
    st.markdown("### 📋 Prompt Settings")

    prompt_type = st.radio(
        "Prompt Type",
        ["Daily Prompts Only", "Friday Fun Prompts Only", "Both"],
        index=2,
        help="Daily prompts focus on practical exercises. Friday prompts are more creative/fun."
    )

    num_daily = num_friday = 0

    if prompt_type in ["Daily Prompts Only", "Both"]:
        num_daily = st.slider(
            "Number of Daily Prompts",
            min_value=1, max_value=31, value=1,
            help="Up to 31 daily prompts per run"
        )

    if prompt_type in ["Friday Fun Prompts Only", "Both"]:
        num_friday = st.slider(
            "Number of Friday Prompts",
            min_value=1, max_value=5, value=1
        )

    st.markdown("---")

    # ── Topic Instructions ────────────────────────────────────────────────────
    st.markdown("### 🎯 Topic Instructions")
    topic_instructions = st.text_area(
        "Custom Instructions / Topic Areas",
        placeholder="e.g., Focus on key AI learning principles, importance of data security, and other similar areas of interest. Target audience: finance team.",
        height=150,
        help="Optional: guide the topics and skill/focus area of generated prompts."
    )

    st.markdown("---")

    # ── Style Options ─────────────────────────────────────────────────────────
    st.markdown("### 🎨 Style Options")
    difficulty = st.select_slider(
        "Difficulty Level",
        options=["Beginner", "Intermediate", "Advanced", "Mixed"],
        value="Mixed"
    )


# ─── AI Caller ────────────────────────────────────────────────────────────────

def _sanitize_key(key: str) -> str:
    """Strip all whitespace, newlines, quotes and invisible characters from an API key."""
    # Remove all whitespace variants (spaces, tabs, newlines, carriage returns)
    k = re.sub(r'\s+', '', key)
    # Strip any surrounding quotes the user may have copy-pasted
    k = k.strip('"\'`')
    # Remove any non-printable / zero-width characters
    k = re.sub(r'[^\x20-\x7E]', '', k)
    return k

def call_ai_agent(user_message: str, system_message: str, agent: str, key: str) -> list:
    """Route to the correct AI agent and return parsed JSON list."""
    if not key or not key.strip():
        raise ValueError(
            "API key is required. Please enter your API key in the sidebar under 🔑 API Key."
        )

    clean_key  = _sanitize_key(key)
    agent_lower = agent.lower()

    if "claude" in agent_lower:
        import anthropic
        # Validate key looks like an Anthropic key before sending
        if not clean_key.startswith("sk-ant-"):
            raise ValueError(
                "Invalid Claude API key format. Claude keys start with 'sk-ant-'. "
                "Please check you copied the full key from console.anthropic.com."
            )
        client = anthropic.Anthropic(api_key=clean_key)
        resp = client.messages.create(
            model="claude-opus-4-6",
            max_tokens=8000,
            system=system_message,
            messages=[{"role": "user", "content": user_message}]
        )
        raw = resp.content[0].text.strip()

    elif "chatgpt" in agent_lower:
        from openai import OpenAI
        if not clean_key.startswith("sk-"):
            raise ValueError(
                "Invalid ChatGPT API key format. OpenAI keys start with 'sk-'. "
                "Please check you copied the full key from platform.openai.com."
            )
        client = OpenAI(api_key=clean_key)
        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user",   "content": user_message}
            ],
            max_tokens=8000,
        )
        raw = resp.choices[0].message.content.strip()

    elif "gemini" in agent_lower:
        import google.generativeai as genai
        genai.configure(api_key=clean_key)
        model = genai.GenerativeModel(
            "gemini-1.5-pro",
            system_instruction=system_message
        )
        resp = model.generate_content(user_message)
        raw = resp.text.strip()

    elif "copilot" in agent_lower:
        # GitHub Copilot requires a GitHub Personal Access Token (ghp_... or github_pat_...)
        # It uses the OpenAI-compatible GitHub Models endpoint
        if not (clean_key.startswith("ghp_") or clean_key.startswith("github_pat_")):
            raise ValueError(
                "GitHub Copilot requires a GitHub Personal Access Token, not an OpenAI key.\n"
                "  1. Go to github.com → Settings → Developer Settings → Personal Access Tokens\n"
                "  2. Click 'Generate new token (classic)'\n"
                "  3. Enable the 'copilot' scope (or use a fine-grained token with Models access)\n"
                "  4. Your token will start with 'ghp_' or 'github_pat_'"
            )
        from openai import OpenAI
        client = OpenAI(
            api_key=clean_key,
            base_url="https://models.inference.ai.azure.com",
        )
        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user",   "content": user_message}
            ],
            max_tokens=8000,
        )
        raw = resp.choices[0].message.content.strip()

    else:
        raise ValueError(f"Unknown AI agent: {agent}")

    # Strip markdown fences if present
    raw = re.sub(r'^```(?:json)?\s*', '', raw.strip())
    raw = re.sub(r'\s*```$', '', raw)

    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        raise ValueError(
            f"The AI returned a response that could not be parsed as JSON.\n"
            f"JSON error: {e}\n"
            f"Response preview: {raw[:300]}"
        )


# ─── Prompt Builders ──────────────────────────────────────────────────────────

def build_system_prompt(agent: str, industry: str) -> str:
    return (
        f"You are an expert AI trainer specializing in {agent} adoption for enterprise practitioners "
        f"in the {industry} industry. "
        f"Create practical, engaging prompts that help employees learn to use {agent} effectively "
        f"in their daily {industry} workflows. "
        "Every prompt must be immediately actionable, realistic, and demonstrate clear business value. "
        "Tailor all scenarios, terminology, and examples specifically to the industry provided. "
        "Return ONLY valid JSON — no markdown fences, no explanation, no preamble."
    )

def build_daily_request(num: int, start: int, topic: str, difficulty: str,
                        agent: str, industry: str) -> str:
    end = start + num - 1
    topic_sec = f"\n\nCustom Instructions / Areas of Interest:\n{topic}" if topic.strip() else ""
    return f"""Generate {num} Daily {agent} AI Prompts for AI adoption training in the {industry} industry.{topic_sec}

Difficulty: {difficulty}
AI Tool: {agent}
Industry: {industry}

Return a JSON array of exactly {num} objects with these exact keys:
- "prompt_id": "DP{start:02d}" through "DP{end:02d}" (sequential number only, e.g. DP01, DP02)
- "ai_prompt": full {agent} prompt (2-4 sentences, specific and actionable, relevant to {industry})
- "email_message": team email/Teams message contextualising the exercise for {industry} professionals (3-5 sentences)
- "learning_objective": what practitioners will learn (1-2 sentences)
- "demonstrated_ai_capability": e.g. Summarization, Draft Generation, Data Analysis, Research
- "test_response": realistic example {agent} output specific to {industry} (4-8 sentences)
- "attachment_required": true or false
- "attachment_description": file description if attachment_required (e.g. "Q3 sales report CSV"), else empty string
- "attachment_filename": suggested filename with extension if attachment_required (e.g. "Q3_Sales_Report.docx", "Patient_Intake_Form.csv"), else empty string
- "attachment_content": if attachment_required, generate the FULL realistic sample file content as plain text that a practitioner can actually use to perform the exercise (minimum 150 words of realistic {industry} data/content). If not required, empty string.

Prompts should progressively build skills and feel like real {industry} work tasks."""

def build_friday_request(num: int, start: int, topic: str, difficulty: str,
                         agent: str, industry: str) -> str:
    end = start + num - 1
    topic_sec = f"\n\nCustom Instructions / Areas of Interest:\n{topic}" if topic.strip() else ""
    return f"""Generate {num} Friday Fun {agent} AI Prompts for AI adoption training in the {industry} industry.{topic_sec}

Difficulty: {difficulty}
AI Tool: {agent}
Industry: {industry}

Friday prompts should be engaging, slightly playful, yet professionally relevant to {industry} — sparking curiosity about AI.

Return a JSON array of exactly {num} objects with these exact keys:
- "prompt_id": "FP{start:02d}" through "FP{end:02d}" (sequential number only, e.g. FP01, FP02)
- "ai_prompt": engaging {agent} prompt tailored to {industry} (2-4 sentences)
- "email_message": upbeat team message introducing the exercise (3-5 sentences, casual but professional, {industry}-relevant)
- "learning_objective": what practitioners will learn (1-2 sentences)
- "demonstrated_ai_capability": e.g. Creative Writing, Brainstorming, Scenario Planning
- "test_response": engaging realistic example {agent} output for {industry} (4-8 sentences)
- "attachment_required": true or false
- "attachment_description": file description if attachment_required (e.g. "Team brainstorm notes"), else empty string
- "attachment_filename": suggested filename with extension if attachment_required (e.g. "Brainstorm_Notes.docx"), else empty string
- "attachment_content": if attachment_required, generate the FULL realistic sample file content as plain text that a practitioner can actually use to perform the exercise (minimum 150 words of realistic {industry} content). If not required, empty string."""


# ─── DOCX Builder ─────────────────────────────────────────────────────────────

def generate_attachment_docx(prompt: dict) -> bytes | None:
    """Generate a real .docx sample attachment from the AI-produced attachment_content."""
    content = prompt.get("attachment_content", "").strip()
    if not content:
        return None

    filename    = prompt.get("attachment_filename", "Attachment.docx")
    description = prompt.get("attachment_description", "Sample attachment")
    pid         = prompt.get("prompt_id", "")

    doc = Document()
    sec = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, attr, Inches(1.0))

    # Title
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    tr = title.add_run(f"{pid} — {description}")
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    doc.add_paragraph()  # spacer

    # Note banner
    note = doc.add_paragraph()
    nr = note.add_run("📋 SAMPLE FILE FOR AI EXERCISE  —  Use this document as input for your prompt.")
    nr.bold = True
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(0x7a, 0x5c, 0x00)
    note.paragraph_format.space_after = Pt(8)

    # Body content — split on newlines, detect headings/bullets
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        import re as _re
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

        # Heading-like lines (ALL CAPS or ends with :)
        if stripped.isupper() or (stripped.endswith(":") and len(stripped) < 60):
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
        # Bullet
        elif stripped.startswith("- ") or stripped.startswith("• "):
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run("• " + stripped[2:])
            run.font.size = Pt(10)
        # Numbered list
        elif _re.match(r"^\d+\.\s", stripped):
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(stripped)
            run.font.size = Pt(10)
        else:
            # Handle **bold** spans
            parts = _re.split(r"(\*\*.+?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r2 = p.add_run(part[2:-2])
                    r2.bold = True
                    r2.font.size = Pt(10)
                elif part:
                    r2 = p.add_run(part)
                    r2.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()



def _make_run(text: str, bold: bool = False) -> OxmlElement:
    """Create a <w:r> with minorHAnsi font, matching template style."""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:asciiTheme'), 'minorHAnsi')
    fonts.set(qn('w:hAnsiTheme'), 'minorHAnsi')
    rPr.append(fonts)
    if bold:
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:bCs'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    return r

def _set_cell_text(cell, text: str, bold: bool = False):
    """Replace paragraphs in a cell, preserving bullets, numbered lists, bold spans and line breaks."""
    import re as _re
    tc = cell._tc
    for p in list(tc.findall(qn('w:p'))):
        tc.remove(p)

    lines = text.split('\n') if text else ['']
    for line in lines:
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)

        stripped = line.strip()
        is_bullet   = stripped.startswith('- ') or stripped.startswith('• ')
        is_numbered = bool(_re.match(r'^\d+\.\s', stripped))

        if is_bullet or is_numbered:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '360')
            ind.set(qn('w:hanging'), '360')
            pPr.append(ind)
            display_line = ('• ' + stripped[2:]) if is_bullet else stripped
        else:
            display_line = line

        rPr_pPr = OxmlElement('w:rPr')
        fonts = OxmlElement('w:rFonts')
        fonts.set(qn('w:asciiTheme'), 'minorHAnsi')
        fonts.set(qn('w:hAnsiTheme'), 'minorHAnsi')
        rPr_pPr.append(fonts)
        if bold:
            rPr_pPr.append(OxmlElement('w:b'))
            rPr_pPr.append(OxmlElement('w:bCs'))
        pPr.append(rPr_pPr)
        p.append(pPr)

        # Split on **bold** spans and render each part
        parts = _re.split(r'(\*\*.+?\*\*)', display_line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.append(_make_run(part[2:-2], bold=True))
            elif part:
                p.append(_make_run(part, bold=bold))

        tc.append(p)

def _get_cell_text(cell) -> str:
    return ''.join(t.text or '' for t in cell._tc.iter(qn('w:t'))).strip()

def _fix_template_ref(doc: Document):
    """Remove broken external .dotx template reference that causes corruption warnings."""
    import re as _re
    # Clear settings.xml.rels of external template attachment
    rels_part = None
    for rel in doc.part.rels.values():
        if 'settings' in rel.reltype.lower() or 'settings' in str(rel.target_ref).lower():
            pass
    # Direct XML fix on settings
    settings_part = doc.settings.element
    for child in list(settings_part):
        if 'attachedTemplate' in child.tag:
            settings_part.remove(child)

def _unique_cells(row):
    """Return only unique cell objects in a row (skips merged duplicates)."""
    seen, result = set(), []
    for c in row.cells:
        if id(c._tc) not in seen:
            seen.add(id(c._tc))
            result.append(c)
    return result

def _update_header_filename(doc: Document, filename: str):
    """Update the header text node to match the output filename."""
    for rel in doc.part.rels.values():
        if 'header' in rel.reltype.lower():
            hdr_part = rel.target_part
            for t in hdr_part._element.iter(qn('w:t')):
                if t.text and t.text.strip():
                    t.text = filename
                    return

def create_prompt_docx(data: dict, is_friday: bool = False,
                       agent: str = "AI", industry: str = "",
                       attachment_bytes: bytes = None,
                       attachment_name: str = None) -> bytes:
    """Fill the real Word template with generated content, matching layout exactly.
    If attachment_bytes is provided, appends the file as an embedded attachment paragraph."""
    template_name = "Friday_Prompt_Template.docx" if is_friday else "Daily_Prompt_Template.docx"
    template_path = TEMPLATES_DIR / template_name

    if not template_path.exists():
        raise FileNotFoundError(
            f"Template not found: {template_path}\n"
            "Please place Daily_Prompt_Template.docx and Friday_Prompt_Template.docx "
            "in the 'templates/' folder next to app.py."
        )

    doc = Document(str(template_path))
    _fix_template_ref(doc)

    # Build file label: "DP01 - Topic" or "FP01 - Topic"
    pid      = data.get('prompt_id', 'DP01')
    topic    = data.get('demonstrated_ai_capability', 'Prompt').strip()
    file_label = f"{pid} - Prompt Topic - {topic}"

    # Update Word header to match filename
    _update_header_filename(doc, file_label)

    # Build attachment text
    att = data.get('attachment_description', '') if data.get('attachment_required') else 'None required'
    data['attachment_text'] = att

    label_to_key = {
        'E-Mail Message':             'email_message',
        'Learning Objective':         'learning_objective',
        'Demonstrated AI Capability': 'demonstrated_ai_capability',
        'Test Response':              'test_response',
        'Attachment (if required)':   'attachment_text',
    }

    tbl = doc.tables[0]
    for row in tbl.rows:
        ucells = _unique_cells(row)
        if not ucells:
            continue

        left_text = _get_cell_text(ucells[0])

        # Header row (DP##-__ / FP##-__):
        # - cell[0] = prompt ID label → inject prompt_id only
        # - cell[1] = "AI Prompt" label → leave as-is (it's a descriptor)
        # - cell[2] = empty content area → inject ai_prompt here
        if 'DP##' in left_text or 'FP##' in left_text:
            _set_cell_text(ucells[0], pid, bold=True)
            # cell[1] is the "AI Prompt" label - DO NOT TOUCH
            # cell[2] is the empty content area
            if len(ucells) > 2:
                _set_cell_text(ucells[2], data.get('ai_prompt', ''))
            continue

        # Content rows: left cell = descriptor label (DO NOT TOUCH)
        # right cell = empty content area → inject content
        for label, key in label_to_key.items():
            if label in left_text and len(ucells) > 1:
                _set_cell_text(ucells[-1], data.get(key, ''))
                break

    # ── Append the generated attachment content after a page break ──
    if attachment_bytes and attachment_name:
        import copy

        # Insert a proper page break via XML on the sectPr paragraph
        pb = doc.add_paragraph()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        pb.add_run()._r.append(br)

        # Section header banner
        hdr = doc.add_paragraph()
        hdr.paragraph_format.space_after = Pt(4)
        hr_run = hdr.add_run(f"SAMPLE ATTACHMENT — {attachment_name}")
        hr_run.bold = True
        hr_run.font.size = Pt(13)
        hr_run.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

        note_p = doc.add_paragraph()
        note_p.paragraph_format.space_after = Pt(10)
        nr = note_p.add_run(
            "This sample file was generated for your AI prompt exercise. "
            "Use it as the input document when running the prompt on the previous page."
        )
        nr.font.size = Pt(9)
        nr.italic = True
        nr.font.color.rgb = RGBColor(0x7a, 0x5c, 0x00)

        # Merge attachment docx body elements into this document
        try:
            att_doc = Document(io.BytesIO(attachment_bytes))
            # sectPr is the last child of body — insert before it
            sect_pr = doc.element.body[-1]
            for element in att_doc.element.body:
                tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
                if tag in ("p", "tbl"):
                    doc.element.body.insert(
                        list(doc.element.body).index(sect_pr),
                        copy.deepcopy(element)
                    )
        except Exception as _e:
            fb = doc.add_paragraph()
            fb.add_run(f"[Could not merge attachment: {_e}]").font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def create_zip(prompts, is_friday_list, agent, industry) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for p, f in zip(prompts, is_friday_list):
            pid      = p.get('prompt_id', 'P01')
            topic    = p.get('demonstrated_ai_capability', 'Prompt').strip()
            has_att  = bool(p.get('attachment_required'))

            # Generate attachment bytes if needed
            att_bytes = generate_attachment_docx(p) if has_att else None
            att_name  = p.get('attachment_filename', 'Attachment.docx') if has_att else None
            if att_name and not att_name.endswith('.docx'):
                att_name = att_name.rsplit('.', 1)[0] + '.docx'

            # Prompt docx (with attachment appended at end)
            prompt_fname = f"{pid} - {topic}.docx"
            zf.writestr(
                prompt_fname,
                create_prompt_docx(p, f, agent, industry,
                                   attachment_bytes=att_bytes,
                                   attachment_name=att_name)
            )

            # Also include the standalone attachment file in the zip
            if att_bytes and att_name:
                # Name: "DP01 - Sample Attachment - Brief Description.docx"
                short_desc = p.get("attachment_description", "Sample").strip()
                # Truncate description to keep filename reasonable
                short_desc = short_desc[:50].rstrip() if len(short_desc) > 50 else short_desc
                safe_desc = re.sub(r'[\\/:*?"<>|]', '-', short_desc)
                att_zip_name = f"{pid} - Sample Attachment - {safe_desc}.docx"
                zf.writestr(att_zip_name, att_bytes)

    buf.seek(0)
    return buf.getvalue()


# ─── Main UI ──────────────────────────────────────────────────────────────────

total = num_daily + num_friday
m1, m2, m3 = st.columns(3)
m1.metric("Daily Prompts",     num_daily)
m2.metric("Friday Prompts",    num_friday)
m3.metric("Total to Generate", total)

st.markdown("---")

if total > 0:
    st.markdown(
        f"**Generating for:** 🤖 `{ai_agent}` &nbsp;|&nbsp; "
        f"🏭 `{industry}` &nbsp;|&nbsp; "
        f"📊 Difficulty: `{difficulty}`",
        unsafe_allow_html=True
    )
    st.markdown("")

if st.button("🚀 Generate Prompts", disabled=(total == 0)):

    if not ai_agent:
        st.error("⚠️ Please select an AI Agent in the sidebar before generating.")
        st.stop()

    if not industry:
        st.error("⚠️ Please select an Industry in the sidebar before generating.")
        st.stop()

    if not api_key or not api_key.strip():
        st.error("⚠️ Please enter your API key in the sidebar under 🔑 API Key before generating.")
        st.stop()

    all_prompts, all_friday = [], []
    bar    = st.progress(0)
    status = st.empty()

    if num_daily > 0:
        status.markdown(
            f'<div class="status-bar">⏳ Generating {num_daily} Daily Prompts with {ai_agent} for {industry}...</div>',
            unsafe_allow_html=True
        )
        try:
            result = call_ai_agent(
                user_message=build_daily_request(num_daily, 1, topic_instructions, difficulty, ai_agent, industry),
                system_message=build_system_prompt(ai_agent, industry),
                agent=ai_agent,
                key=api_key
            )
            if isinstance(result, list):
                all_prompts.extend(result)
                all_friday.extend([False] * len(result))
            bar.progress(0.5 if num_friday > 0 else 1.0)
        except Exception as e:
            st.error(f"Error generating daily prompts: {e}")
            st.stop()

    if num_friday > 0:
        status.markdown(
            f'<div class="status-bar">⏳ Generating {num_friday} Friday Fun Prompts with {ai_agent} for {industry}...</div>',
            unsafe_allow_html=True
        )
        try:
            result = call_ai_agent(
                user_message=build_friday_request(num_friday, 1, topic_instructions, difficulty, ai_agent, industry),
                system_message=build_system_prompt(ai_agent, industry),
                agent=ai_agent,
                key=api_key
            )
            if isinstance(result, list):
                all_prompts.extend(result)
                all_friday.extend([True] * len(result))
            bar.progress(1.0)
        except Exception as e:
            st.error(f"Error generating Friday prompts: {e}")
            st.stop()

    status.markdown(
        f'<div class="status-bar">✅ Generated {len(all_prompts)} prompt(s) with {ai_agent} for {industry}!</div>',
        unsafe_allow_html=True
    )
    st.session_state['generated_prompts'] = all_prompts
    st.session_state['is_friday_list']    = all_friday
    st.session_state['last_agent']        = ai_agent
    st.session_state['last_industry']     = industry


# ─── Results ──────────────────────────────────────────────────────────────────

if st.session_state.get('generated_prompts'):
    prompts       = st.session_state['generated_prompts']
    fri_list      = st.session_state['is_friday_list']
    last_agent    = st.session_state.get('last_agent', ai_agent)
    last_industry = st.session_state.get('last_industry', industry)

    st.markdown("## 📄 Generated Prompts")

    # ── Download All (no wrapping div that causes the blue box) ──────────────
    with st.container(border=True):
        _, dcol, _ = st.columns([1, 2, 1])
        with dcol:
            with st.spinner("Packaging files..."):
                zip_bytes = create_zip(prompts, fri_list, last_agent, last_industry)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            st.download_button(
                label=f"⬇️  Download All Extract Files for All {len(prompts)} Prompts (.zip)",
                data=zip_bytes,
                file_name=f"AI_Prompts_{last_agent.replace(' ','_')}_{last_industry.replace(' ','_')}_{ts}.zip",
                mime="application/zip",
                use_container_width=True
            )
            st.markdown(
                f"<div style='text-align:center;font-size:0.82rem;color:#555;margin-top:0.3rem;'>"
                f"📦 {len(prompts)} DOCX files &nbsp;|&nbsp; "
                f"🤖 {last_agent} &nbsp;|&nbsp; 🏭 {last_industry}</div>",
                unsafe_allow_html=True
            )

    # ── Color-coding legend ───────────────────────────────────────────────────
    st.markdown(
        '<div class="legend-bar">'
        '  <strong>Color Guide:</strong>'
        '  <span class="legend-item"><span class="legend-dot" style="background:#4472c4;border:2px solid #002060;"></span> Daily Prompt</span>'
        '  <span class="legend-item"><span class="legend-dot" style="background:#7030a0;"></span> Friday Fun Prompt</span>'
        '  <span class="legend-item"><span class="legend-dot" style="background:#ffe082;"></span> Attachment Required</span>'
        '  <span class="legend-item"><span class="legend-dot" style="background:#c6efce;"></span> No Attachment</span>'
        '</div>',
        unsafe_allow_html=True
    )

    # ── Prompt expanders ──────────────────────────────────────────────────────
    def fmt(text: str) -> str:
        """Escape HTML then restore newlines and basic markdown for display."""
        import html as _html
        import re as _re
        t = _html.escape(str(text))
        t = _re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', t)
        lines = t.split('\n')
        out = []
        for line in lines:
            stripped = line.strip()
            if stripped.startswith('- ') or stripped.startswith('• '):
                out.append(f'&nbsp;&nbsp;• {stripped[2:]}')
            elif _re.match(r'^\d+\.\ ', stripped):
                out.append(f'&nbsp;&nbsp;{stripped}')
            else:
                out.append(line)
        return '<br>'.join(out)

    for i, (prompt, is_fri) in enumerate(zip(prompts, fri_list)):
        icon     = "🟣" if is_fri else "🔵"
        badge    = "badge-friday" if is_fri else "badge-daily"
        label    = "🎉 Friday Fun" if is_fri else "📋 Daily"
        has_att  = bool(prompt.get('attachment_required'))
        att_desc = prompt.get('attachment_description', 'See instructions')
        pid      = prompt.get('prompt_id', f'P{i+1:02d}')
        topic    = prompt.get('demonstrated_ai_capability', 'Prompt').strip()

        with st.expander(
            f"{icon} {pid} — {topic}",
            expanded=(i == 0)
        ):
            st.markdown(f'<span class="{badge}">{label}</span>', unsafe_allow_html=True)

            cl, cr = st.columns([3, 2])

            with cl:
                for lbl, key in [
                    ("💡 AI Prompt",      "ai_prompt"),
                    ("📧 E-Mail Message",  "email_message"),
                    ("🤖 Test Response",   "test_response")
                ]:
                    st.markdown(f'<div class="field-label">{lbl}</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="field-value">{fmt(prompt.get(key, ""))}</div>', unsafe_allow_html=True)

            with cr:
                for lbl, key in [
                    ("🎯 Learning Objective",        "learning_objective"),
                    ("⚡ Demonstrated AI Capability", "demonstrated_ai_capability")
                ]:
                    st.markdown(f'<div class="field-label">{lbl}</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="field-value">{fmt(prompt.get(key, ""))}</div>', unsafe_allow_html=True)

                # ── Attachment section ────────────────────────────────────────
                if has_att:
                    att_fname   = prompt.get("attachment_filename", "Attachment.docx")
                    att_docx    = generate_attachment_docx(prompt)

                    st.markdown(
                        f'<div class="attachment-box">'
                        f'  <div class="att-title">📎 Attachment Required</div>'
                        f'  <div class="att-desc">📄 {att_desc}</div>'
                        f'</div>',
                        unsafe_allow_html=True
                    )
                    if att_docx:
                        st.download_button(
                            label=f"📥 Download Sample Attachment — {prompt.get('attachment_description', 'Sample')[:45]}...",
                            data=att_docx,
                            file_name=att_fname if att_fname.endswith(".docx") else att_fname.rsplit(".",1)[0] + ".docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"att_dl_{i}",
                            use_container_width=True,
                            help="Download this sample file, then use it as input when running your AI prompt exercise."
                        )
                    else:
                        st.info("ℹ️ No sample content was generated for this attachment.")
                else:
                    st.markdown('<div class="field-label">📎 Attachment</div>', unsafe_allow_html=True)
                    st.markdown('<div class="field-value" style="background:#f0faf3;border-color:#c6efce;">❌ Not required</div>', unsafe_allow_html=True)

                # ── Download this prompt (with attachment appended) ────────
                st.markdown("**⬇️ Download prompt**")
                att_docx_for_embed  = generate_attachment_docx(prompt) if has_att else None
                att_fname_for_embed = prompt.get("attachment_filename", "Attachment.docx") if has_att else None
                if att_fname_for_embed and not att_fname_for_embed.endswith(".docx"):
                    att_fname_for_embed = att_fname_for_embed.rsplit(".", 1)[0] + ".docx"
                st.download_button(
                    label=f"⬇ {pid} - {topic}.docx",
                    data=create_prompt_docx(
                        prompt, is_fri, last_agent, last_industry,
                        attachment_bytes=att_docx_for_embed,
                        attachment_name=att_fname_for_embed
                    ),
                    file_name=f"{pid} - {topic}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_{i}",
                    use_container_width=True
                )
else:
    st.markdown("""
    <div style="text-align:center;padding:4rem 2rem;background:white;border-radius:12px;box-shadow:0 2px 10px rgba(0,0,0,0.07);">
        <div style="font-size:4rem;margin-bottom:1rem;">🤖</div>
        <h3 style="color:#002060;margin-bottom:0.5rem;">Ready to Generate Prompts</h3>
        <p style="color:#666;max-width:500px;margin:0 auto;">
            Configure your settings in the sidebar — select an AI agent, industry, and prompt quantity,
            add optional topic instructions, then click <strong>Generate Prompts</strong>.
        </p>
        <br/>
        <p style="color:#999;font-size:0.85rem;">📋 Up to 31 Daily Prompts &nbsp;•&nbsp; 🎉 Up to 5 Friday Fun Prompts &nbsp;•&nbsp; ⬇️ Bulk ZIP Download</p>
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")
st.markdown("""
<div style="text-align:center;color:#999;font-size:0.8rem;padding:1rem;">
    AI Prompt Generator &nbsp;•&nbsp; Multi-Agent Support &nbsp;•&nbsp;
    Restricted to SSA & Company and authorized users
</div>
""", unsafe_allow_html=True)
